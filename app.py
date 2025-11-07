import streamlit as st
from PIL import Image
import io
import base64
import pandas as pd
from docx import Document
from docx.shared import Inches
# Importa la librería de Vertex AI
import vertexai
from vertexai.generative_models import GenerativeModel, Part, Image as VertexImage, GenerationConfig
import json
import random # Necesario para la clave aleatoria

# --- IMPORTACIÓN CLAVE ---
# Importamos 'crear_grafico' (el renderizador) y no 'generar_grafico_desde_texto' (el que llama al LLM)
try:
    from graficos_plugins import crear_grafico
    GRAFICOS_DISPONIBLES = True
except ImportError:
    st.error("Advertencia: No se encontró el archivo 'graficos_plugins.py'. La previsualización de gráficos no funcionará.")
    GRAFICOS_DISPONIBLES = False
    # Definir una función placeholder si falla la importación
    def crear_grafico(*args, **kwargs):
        return None

# --- Configuración de Google Cloud (hacer al inicio) ---
# Descomenta esta línea y configúrala con tu proyecto y región
# vertexai.init(project="TU_PROYECTO_GCP", location="TU_REGION")

# --- 1. FUNCIÓN DEL GENERADOR (ACTUALIZADA) ---
# --- 1. FUNCIÓN DEL GENERADOR (ACTUALIZADA Y MEJORADA) ---

# --- 1. FUNCIÓN DEL GENERADOR (ACTUALIZADA Y MEJORADA) ---
def generar_item_llm(imagen_cargada, taxonomia_dict, contexto_adicional, feedback_auditor=""):
    """
    GENERADOR: Genera el ítem, donde el enunciado Y/O las opciones pueden ser imágenes/tablas.
    (Versión mejorada con limpieza de JSON y lógica de gráficos avanzada)
    """
    
    # --- Configuración del Modelo ---
    model = GenerativeModel("gemini-2.5-flash-lite") 
    
    # --- Procesamiento de Imagen ---
    img_pil = Image.open(imagen_cargada)
    buffered = io.BytesIO()
    img_pil.save(buffered, format="PNG")
    img_bytes = buffered.getvalue()
    vertex_img = VertexImage.from_bytes(img_bytes)

    # --- Preparación de variables del Prompt ---
    taxonomia_texto = "\n".join([f"* {k}: {v}" for k, v in taxonomia_dict.items()])
    clave_aleatoria = random.choice(['A', 'B', 'C', 'D'])

    seccion_feedback = ""
    if feedback_auditor:
        seccion_feedback = f"""
        --- RETROALIMENTACIÓN DE AUDITORÍA (Error a corregir) ---
        El intento anterior fue rechazado. DEBES corregir los siguientes errores:
        {feedback_auditor}
        --- VUELVE A GENERAR EL ÍTEM CORRIGIENDO ESTO ---
        """

    # --- 4. Diseño del Prompt (Generador) - ¡CON LÓGICA DE GRÁFICOS INTEGRADA! ---
    prompt_texto = f"""
    Eres un psicómetra experto en "Shells Cognitivos". Tu tarea es crear un ítem espejo basado en la imagen adjunta, alineado con la taxonomía y el contexto.
    DEBES devolver un JSON válido.

    {seccion_feedback}

    **Shell Cognitivo (Pregunta Original):**
    Analiza la estructura lógica y la "Tarea Cognitiva" de la pregunta en la IMAGEN ADJUNTA.
    - Si la pregunta original usa una tabla o gráfico, tu ítem espejo también debería usar uno.
    - **¡IMPORTANTE!** Si las *opciones de respuesta* en la imagen original son gráficas o tablas, debes replicar esa estructura para las opciones del ítem espejo.

    **Taxonomía Requerida (Tu Guía):**
    {taxonomia_texto}
    
    **Contexto Adicional del Usuario (Tema del ítem nuevo):**
    {contexto_adicional}

    --- ANÁLISIS COGNITIVO OBLIGATORIO (Tu paso 1) ---
    Basado en la taxonomía (Evidencia, Afirmación, Competencia), define la Tarea Cognitiva exacta que el ítem espejo debe evaluar.
    
    --- CONSTRUCCIÓN DEL ÍTEM (Tu paso 2) ---
    Basado en tu análisis, construye el ítem.
    - ENUNCIADO: Debe ser claro y **NO** usar jerarquías ("más", "mejor", "principalmente").
    - CLAVE: La respuesta correcta DEBE ser la opción **{clave_aleatoria}**.
    - DISTRACTORES: Plausibles, basados en errores comunes de la Tarea Cognitiva.
    
    
    --- INSTRUCCIONES DE SALIDA PARA GRÁFICO (ENUNCIADO Y OPCIONES) ---
    Tanto el enunciado ("descripcion_grafico_enunciado") como CADA opción ("descripcion_grafico") 
    pueden contener gráficos.

    Si el elemento (enunciado u opción) NO necesita un gráfico, usa "NO" y [].
    Si SÍ necesita un gráfico, usa "SÍ" y proporciona una LISTA DE OBJETOS JSON VÁLIDOS 
    (incluso si es un solo gráfico).

    Cada objeto JSON en la lista DEBE contener: "tipo_elemento", "datos", "configuracion" y "descripcion".

    1. Para "tipo_elemento", elige UNO de la siguiente lista: 
       grafico_barras_verticales, grafico_circular, tabla, construccion_geometrica, 
       diagrama_arbol, flujograma, pictograma, scatter_plot, line_plot, 
       histogram, box_plot, otro_tipo.
       
    2. Para "descripcion", proporciona un texto en lenguaje natural que resuma el gráfico 
       para validación.

    3. LÓGICA CONDICIONAL PARA EL CAMPO "datos":
       - Si eliges un tipo de la lista (QUE NO SEA "otro_tipo"): 
         El campo "datos" debe ser un objeto con la información estructurada.
         (Ej: {{"columnas": ["X", "Y"], "filas": [[1, 2]]}})
       - Si eliges "otro_tipo" (para diagramas, geometrías, etc.):
         El campo "datos" debe ser un objeto con una clave "descripcion_natural".
         (Ej: {{"descripcion_natural": "Un diagrama de un circuito en serie con una batería de 9V y tres resistencias..."}})

    --- FORMATO DE SALIDA OBLIGATORIO (JSON VÁLIDO) ---
    Responde ÚNICAMENTE con el objeto JSON. No incluyas ```json.
    {{
      "pregunta_espejo": "Texto completo del enunciado/stem...",
      "clave": "{clave_aleatoria}",
      "descripcion_imagen_original": "Descripción de la imagen que el usuario subió...",
      "justificacion_clave": "Razón por la que la clave es correcta...",
      
      "grafico_necesario_enunciado": "SÍ",
      "descripcion_grafico_enunciado": [
        {{
          "tipo_elemento": "tabla",
          "datos": {{ "columnas": ["País", "Capital"], "filas": [["Colombia", "Bogotá"]] }},
          "configuracion": {{ "titulo": "Capitales" }},
          "descripcion": "Una tabla simple de países y capitales."
        }}
      ],
      
      "opciones": {{
        "A": {{
          "texto": "Ver gráfico A",
          "grafico_necesario": "SÍ",
          "descripcion_grafico": [
            {{
              "tipo_elemento": "otro_tipo",
              "datos": {{ "descripcion_natural": "Un diagrama de un circuito eléctrico simple en serie..." }},
              "configuracion": {{ "titulo": "Circuito en Serie" }},
              "descripcion": "Diagrama de un circuito en serie."
            }}
          ]
        }},
        "B": {{
          "texto": "Texto de la Opción B (sin gráfico)",
          "grafico_necesario": "NO",
          "descripcion_grafico": []
        }},
        "C": {{
          "texto": "Texto de la Opción C",
          "grafico_necesario": "NO",
          "descripcion_grafico": []
        }},
        "D": {{
          "texto": "Texto de la Opción D",
          "grafico_necesario": "NO",
          "descripcion_grafico": []
        }}
      }},
      
      "justificaciones_distractores": [
        {{ "opcion": "A", "justificacion": "Justificación para A..." }},
        {{ "opcion": "B", "justificacion": "Justificación para B..." }},
        {{ "opcion": "C", "justificacion": "Justificación para C..." }},
        {{ "opcion": "D", "justificacion": "Justificación para D..." }}
      ]
    }}
    """

    config_generacion = GenerationConfig(
        response_mime_type="application/json"
    )

    try:
        # --- 1. LLAMADA A LA API ---
        response = model.generate_content(
            [vertex_img, prompt_texto], 
            generation_config=config_generacion
        )
        
        raw_text = response.text
        
        # --- 2. MEJORA: LIMPIEZA DE JSON ---
        # (Esto resuelve el error de 'Error al parsear el JSON final')
        try:
            # Encuentra el primer { y el último } para eliminar texto extra
            start_index = raw_text.find('{')
            end_index = raw_text.rfind('}') + 1
            
            if start_index == -1 or end_index == 0:
                raise ValueError("No se encontraron los delimitadores JSON '{' o '}'.")

            # Extrae solo el JSON
            json_str = raw_text[start_index:end_index]
            
            # Valida que es un JSON antes de devolver
            json.loads(json_str) 
            return json_str
        
        except (ValueError, json.JSONDecodeError) as json_e:
            st.error(f"Error al limpiar/parsear la respuesta del Generador: {json_e}")
            st.error(f"Respuesta cruda recibida (esto puede ayudar a depurar): {raw_text}")
            return None
        # --- FIN DE LA MEJORA DE LIMPIEZA ---

    except Exception as e:
        st.error(f"Error al contactar Vertex AI (Generador): {e}")
        return None

# --- 2. FUNCIÓN DEL AUDITOR (ACTUALIZADA CON LIMPIEZA DE JSON) ---
def auditar_item_llm(item_json_texto, taxonomia_dict):
    """
    AUDITOR: Audita el ítem Y la coherencia de los gráficos (enunciado Y opciones).
    """
    
    # Modelo de Gemini (corregido al que usas)
    model = GenerativeModel("gemini-2.5-flash-lite")
    taxonomia_texto = "\n".join([f"* {k}: {v}" for k, v in taxonomia_dict.items()])

    prompt_auditor = f"""
    Eres un auditor psicométrico experto y riguroso. Tu tarea es auditar el siguiente ítem (en JSON)
    contra la taxonomía y las reglas de estilo.
    
    **Taxonomía de Referencia (ObligatorIA):**
    {taxonomia_texto}

    **Ítem Generado (JSON a Auditar):**
    {item_json_texto}

    --- CRITERIOS DE AUDITORÍA (Evalúa uno por uno) ---
    1.  **Alineación con Taxonomía:** ¿El ítem evalúa CLARAMENTE la Evidencia, Afirmación y Competencia?
    2.  **Estilo del Enunciado (No Jerarquización):** ¿El enunciado usa palabras prohibidas como "más", "mejor", "principalmente"? (RECHAZO automático).
    3.  **Calidad de Distractores:** ¿Las justificaciones de los distractores explican el *error* (ej. "El estudiante podría...")?
    4.  **Clave y Opciones:** ¿Hay 4 opciones? ¿La clave coincide con una opción?
    5.  **Coherencia de Gráficos (¡ACTUALIZADO!):** - ¿Es coherente el "grafico_necesario_enunciado" con la pregunta?
        - ¿Son coherentes los "grafico_necesario" DENTRO de cada opción?
        - Si un gráfico existe, ¿es un JSON válido?

    --- FORMATO DE SALIDA OBLIGATORIO (JSON VÁLIDO) ---
    Devuelve tu auditoría como un único objeto JSON. No uses ```json.
    {{
      "criterios": [
        {{ "criterio": "1. Alineación con Taxonomía", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "2. Estilo (No Jerarquización)", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "3. Calidad de Distractores", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "4. Clave y Opciones", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "5. Coherencia de Gráficos", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }}
      ],
      "dictamen_final": "✅ CUMPLE" o "❌ RECHAZADO",
      "observaciones_finales": "Si es RECHAZADO, explica aquí CLARAMENTE qué debe corregir el generador. (Ej: 'El enunciado usa la palabra 'principalmente'. O 'El gráfico de la opción C es SÍ pero no se proporcionó JSON.')"
    }}
    """
    
    config_generacion = GenerationConfig(
        response_mime_type="application/json"
    )

    try:
        response = model.generate_content(
            prompt_auditor, 
            generation_config=config_generacion
        )
        
        raw_text = response.text
        
        # --- INICIO DE LA MEJORA: LIMPIEZA DE JSON (AUDITOR) ---
        try:
            # Encuentra el primer { y el último } para eliminar texto extra
            start_index = raw_text.find('{')
            end_index = raw_text.rfind('}') + 1
            
            if start_index == -1 or end_index == 0:
                raise ValueError("No se encontraron los delimitadores JSON '{' o '}'.")

            # Extrae solo el JSON
            json_str = raw_text[start_index:end_index]
            
            # Valida que es un JSON antes de devolver
            json.loads(json_str) 
            return json_str
        
        except (ValueError, json.JSONDecodeError) as json_e:
            st.error(f"Error al limpiar/parsear la respuesta del Auditor: {json_e}")
            st.error(f"Respuesta cruda recibida (esto puede ayudar a depurar): {raw_text}")
            return None
        # --- FIN DE LA MEJORA DE LIMPIEZA ---

    except Exception as e:
        st.error(f"Error al contactar Vertex AI (Auditor): {e}")
        return None

# --- 3. FUNCIONES DE EXPORTACIÓN (ACTUALIZADAS) ---

def crear_excel(datos_generados):
    data_rows = []
    data_rows.append({"Componente": "Pregunta Espejo", "Contenido": datos_generados.get("pregunta_espejo", "")})
    
    # Añadir info del gráfico del enunciado
    data_rows.append({"Componente": "Gráfico Enunciado", "Contenido": datos_generados.get("grafico_necesario_enunciado", "NO")})
    grafico_json_str = json.dumps(datos_generados.get("descripcion_grafico_enunciado", []), indent=2)
    data_rows.append({"Componente": "Datos Gráfico Enunciado (JSON)", "Contenido": grafico_json_str})

    opciones = datos_generados.get("opciones", {})
    for letra in ["A", "B", "C", "D"]:
        opcion_obj = opciones.get(letra, {})
        # Añadir texto de la opción
        data_rows.append({"Componente": f"Opción {letra} - Texto", "Contenido": opcion_obj.get("texto", "")})
        # Añadir info del gráfico de la opción
        data_rows.append({"Componente": f"Opción {letra} - Gráfico", "Contenido": opcion_obj.get("grafico_necesario", "NO")})
        grafico_json_str = json.dumps(opcion_obj.get("descripcion_grafico", []), indent=2)
        data_rows.append({"Componente": f"Opción {letra} - Datos Gráfico (JSON)", "Contenido": grafico_json_str})

    data_rows.append({"Componente": "Clave", "Contenido": datos_generados.get("clave", "")})
    data_rows.append({"Componente": "Justificación Clave", "Contenido": datos_generados.get("justificacion_clave", "")})
    justificaciones = datos_generados.get("justificaciones_distractores", [])
    for just in justificaciones:
        data_rows.append({"Componente": f"Justificación {just.get('opcion')}", "Contenido": just.get('justificacion')})
    
    df = pd.DataFrame(data_rows)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Item Generado')
    return output.getvalue()

def crear_word(datos_generados):
    document = Document()
    document.add_heading('Ítem Espejo Generado', level=1)
    
    document.add_heading('Pregunta Espejo (Enunciado)', level=2)
    document.add_paragraph(datos_generados.get("pregunta_espejo", "N/A"))
    
    if datos_generados.get("grafico_necesario_enunciado") == "SÍ":
        document.add_heading('Datos del Gráfico (Enunciado)', level=3)
        grafico_json_str = json.dumps(datos_generados.get("descripcion_grafico_enunciado", []), indent=2)
        document.add_paragraph(grafico_json_str)

    document.add_heading('Opciones', level=2)
    opciones = datos_generados.get("opciones", {})
    for letra in ["A", "B", "C", "D"]:
        opcion_obj = opciones.get(letra, {})
        document.add_heading(f"Opción {letra}", level=3)
        document.add_paragraph(opcion_obj.get("texto", "N/A"))
        
        if opcion_obj.get("grafico_necesario") == "SÍ":
            document.add_heading(f'Datos del Gráfico (Opción {letra})', level=4)
            grafico_json_str = json.dumps(opcion_obj.get("descripcion_grafico", []), indent=2)
            document.add_paragraph(grafico_json_str)

    document.add_heading('Clave', level=2)
    document.add_paragraph(datos_generados.get('clave', 'N/A'))
    
    document.add_heading('Justificaciones', level=2)
    document.add_paragraph(f"**Justificación de la Clave:** {datos_generados.get('justificacion_clave', 'N/A')}")
    document.add_heading('Justificaciones de Distractores', level=3)
    justificaciones = datos_generados.get("justificaciones_distractores", [])
    for just in justificaciones:
        if just.get('opcion') != datos_generados.get('clave'):
            document.add_paragraph(f"**Justificación {just.get('opcion')}:** {just.get('justificacion')}")
    
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

# --- 4. INTERFAZ DE STREAMLIT (UI) ---

st.set_page_config(layout="wide")
st.title("🤖 Generador de Ítemes (con Auditoría de IA)")

# --- Columnas para la entrada ---
col1, col2 = st.columns(2)

with col1:
    st.header("1. Cargar Ítem Original")
    imagen_subida = st.file_uploader(
        "Sube el pantallazo de la pregunta", 
        type=["png", "jpg", "jpeg"]
    )
    
    if imagen_subida:
        st.image(imagen_subida, caption="Ítem cargado", use_container_width=True)

# --- COLUMNA 2 (Lógica de Filtros Bifurcada y CORREGIDA) ---
with col2:
    st.header("2. Configurar Generación")
    
    excel_file = st.file_uploader("Cargar Excel de Taxonomía (un solo .xlsx)", type=['xlsx'])
    
    grado_sel, area_sel, comp1_sel, comp2_sel, ref_sel, competen_sel, afirm_sel, evid_sel = (None,) * 8
    
    if excel_file is not None:
        try:
            if 'df1' not in st.session_state or 'df2' not in st.session_state:
                data = pd.read_excel(excel_file, sheet_name=None)
                sheet_names = list(data.keys())
                if len(sheet_names) < 2:
                    st.error("Error: El archivo Excel debe tener al menos dos hojas.")
                    excel_file = None
                else:
                    st.session_state.df1 = data[sheet_names[0]]
                    st.session_state.df2 = data[sheet_names[1]]
                    st.success(f"Éxito: Cargadas hojas '{sheet_names[0]}' y '{sheet_names[1]}'.")
            
            if 'df1' in st.session_state:
                df1 = st.session_state.df1
                df2 = st.session_state.df2

                # --- Filtros Comunes ---
                grados = df1['Grado'].unique()
                grado_sel = st.selectbox("Grado", options=grados)
                
                df_grado_h1 = df1[df1['Grado'] == grado_sel]
                areas = df_grado_h1['Área'].unique() # Con tilde
                area_sel = st.selectbox("Área", options=areas) # Con tilde

                # --- Cascada 1: (Hoja 1 - Estructura) ---
                st.subheader("Taxonomía (Hoja 1 - Estructura)")
                df_area_h1 = df_grado_h1[df_grado_h1['Área'] == area_sel]
                # --- CORRECCIÓN DE BUG: Usa 'Componente' ---
                componentes1 = df_area_h1['Componente1'].unique() 
                comp1_sel = st.selectbox("Componente (Estructura)", options=componentes1) 

                df_comp1 = df_area_h1[df_area_h1['Componente1'] == comp1_sel]
                competencias = df_comp1['Competencia'].unique()
                competen_sel = st.selectbox("Competencia", options=competencias)

                df_competencia = df_comp1[df_comp1['Competencia'] == competen_sel]
                
                if area_sel == 'Ciencias Naturales': 
                    df_afirmacion_base = df_competencia[df_competencia['Componente1'] == comp1_sel]
                else:
                    df_afirmacion_base = df_competencia
                    
                afirmaciones = df_afirmacion_base['Afirmación'].unique()
                afirm_sel = st.selectbox("Afirmación", options=afirmaciones)

                df_afirmacion = df_afirmacion_base[df_afirmacion_base['Afirmación'] == afirm_sel]
                evidencias = df_afirmacion['Evidencia'].unique()
                evid_sel = st.selectbox("Evidencia", options=evidencias)

                # --- Cascada 2: (Hoja 2 - Temática) ---
                st.subheader("Taxonomía (Hoja 2 - Temática)")
                df_area_h2 = df2[
                    (df2['Grado'] == grado_sel) & 
                    (df2['Área'] == area_sel) # Con tilde
                ]
                # --- CORRECCIÓN DE BUG: Usa 'Componente' ---
                componentes2 = df_area_h2['Componente2'].unique()
                comp2_sel = st.selectbox("Componente (Temática)", options=componentes2)

                df_comp2 = df_area_h2[df_area_h2['Componente2'] == comp2_sel]
                
                refs = df_comp2['Ref. Temática'].unique() if not df_comp2.empty else ["N/A"] # Con tilde y espacio
                ref_sel = st.selectbox("Ref. Temática", options=refs) # Con tilde y espacio

        except KeyError as e:
            st.error(f"Error de Columna: No se encontró la columna {e}. Revisa las tildes/mayúsculas.")
            if 'df1' in st.session_state: st.error(f"Columnas H1: {list(st.session_state.df1.columns)}")
            if 'df2' in st.session_state: st.error(f"Columnas H2: {list(st.session_state.df2.columns)}")
            excel_file = None
        except Exception as e:
            st.error(f"Error inesperado al procesar el Excel: {e}")
            excel_file = None
    
    info_adicional = st.text_area(
        "Contexto Adicional (Tema para el ítem)",
        height=150,
        placeholder="Ej: 'Usar el tema de la fotosíntesis', 'Basarse en la Revolución Francesa'"
    )

# --- 5. LÓGICA DEL BOTÓN (Bucle Generador-Auditor) ---
st.divider()
if st.button("🚀 Generar Ítem Espejo (con Auditoría)", use_container_width=True, type="primary"):
    
    if imagen_subida is None:
        st.warning("Por favor, sube una imagen primero.")
    elif excel_file is None:
        st.warning("Por favor, carga el archivo Excel de taxonomía.")
    elif evid_sel is None or ref_sel is None:
        st.warning("Completa toda la selección de taxonomía.")
    else:
        taxonomia_seleccionada = {
            "Grado": grado_sel,
            "Área": area_sel,
            "Componente_Estructura": comp1_sel, # Nombre corregido
            "Componente_Tematica": comp2_sel,  # Nombre corregido
            "Ref. Temática": ref_sel,
            "Competencia": competen_sel,
            "Afirmación": afirm_sel,
            "Evidencia": evid_sel
        }
        
        max_intentos = 3
        intento_actual = 0
        feedback_auditor = ""
        item_final_json = None

        with st.status("Iniciando proceso...", expanded=True) as status:
            while intento_actual < max_intentos:
                intento_actual += 1
                
                status.update(label=f"Intento {intento_actual}/{max_intentos}: Generando ítem...")
                item_json_str = generar_item_llm(
                    imagen_subida, 
                    taxonomia_seleccionada,
                    info_adicional,
                    feedback_auditor 
                )
                
                if item_json_str is None:
                    status.update(label=f"Error en la generación (Intento {intento_actual}).", state="error")
                    continue 

                status.update(label=f"Intento {intento_actual}/{max_intentos}: Auditando ítem...")
                audit_json_str = auditar_item_llm(item_json_str, taxonomia_seleccionada)

                if audit_json_str is None:
                    status.update(label=f"Error en la auditoría (Intento {intento_actual}).", state="error")
                    continue 

                try:
                    # --- FIX: Asegurarse de parsear la respuesta del auditor ---
                    audit_data = json.loads(audit_json_str)
                    
                    if audit_data.get("dictamen_final") == "✅ CUMPLE":
                        status.update(label="¡Auditoría Aprobada!", state="complete")
                        item_final_json = item_json_str
                        break 
                    else:
                        feedback_auditor = audit_data.get("observaciones_finales", "Rechazado sin observaciones.")
                        status.update(label=f"Intento {intento_actual} Rechazado. Preparando re-intento...")
                        st.expander(f"Detalles del Rechazo (Intento {intento_actual})").json(audit_data)
                
                except json.JSONDecodeError:
                    st.error(f"Error al leer respuesta JSON del auditor: {audit_json_str}")
                    feedback_auditor = "La respuesta del auditor no fue un JSON válido."

            if item_final_json is None:
                status.update(label=f"No se pudo generar un ítem de alta calidad después de {max_intentos} intentos.", state="error")
                st.error(f"Último feedback del auditor: {feedback_auditor}")
            
        if item_final_json:
            st.success("¡Ítem generado y auditado con éxito! Puedes editarlo abajo.")
            try:
                # --- FIX: Asegurarse de parsear la respuesta del generador ---
                datos_obj = json.loads(item_final_json)
                st.session_state['resultado_json_obj'] = datos_obj
                
                # --- LÓGICA DE INICIALIZACIÓN (ACTUALIZADA para nuevo JSON) ---
                st.session_state.editable_pregunta = datos_obj.get("pregunta_espejo", "")
                st.session_state.editable_clave = datos_obj.get("clave", "")
                st.session_state.editable_just_clave = datos_obj.get("justificacion_clave", "")

                # Gráfico del Enunciado
                st.session_state.editable_grafico_nec_enunciado = datos_obj.get("grafico_necesario_enunciado", "NO")
                grafico_data_enunciado = datos_obj.get("descripcion_grafico_enunciado", [])
                st.session_state.editable_grafico_json_enunciado = json.dumps(grafico_data_enunciado, indent=2)

                # Opciones (A, B, C, D)
                opciones = datos_obj.get("opciones", {})
                for letra in ["A", "B", "C", "D"]:
                    # --- FIX: Corregir la inicialización de opciones ---
                    # El JSON antiguo era "A": "Texto". El nuevo es "A": {"texto": "..."}
                    opcion_obj = opciones.get(letra, {}) # Obtener el objeto de la opción
                    
                    st.session_state[f"editable_opcion_{letra.lower()}_texto"] = opcion_obj.get("texto", "")
                    st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"] = opcion_obj.get("grafico_necesario", "NO")
                    grafico_data = opcion_obj.get("descripcion_grafico", [])
                    st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"] = json.dumps(grafico_data, indent=2)

                # Justificaciones
                justifs_list = datos_obj.get("justificaciones_distractores", [])
                justifs_map = {j.get('opcion'): j.get('justificacion') for j in justifs_list}
                st.session_state.editable_just_a = justifs_map.get("A", "N/A")
                st.session_state.editable_just_b = justifs_map.get("B", "N/A")
                st.session_state.editable_just_c = justifs_map.get("C", "N/A")
                st.session_state.editable_just_d = justifs_map.get("D", "N/A")
                
                st.session_state.show_editor = True
                
            except json.JSONDecodeError:
                st.error(f"Error al parsear el JSON final: {item_final_json}")
                st.session_state.show_editor = False

# --- 6. EDITOR DE ÍTEMS Y DESCARGA (ACTUALIZADO) ---
if 'show_editor' in st.session_state and st.session_state.show_editor:
    st.divider()
    st.header("3. Edita el Ítem Generado")
    
    # --- ENUNCIADO Y GRÁFICO DEL ENUNCIADO ---
    st.subheader("Enunciado")
    st.text_area("Texto del Enunciado", key="editable_pregunta", height=150)
    st.selectbox(
        "¿Enunciado necesita un gráfico/tabla?", 
        options=["NO", "SÍ"], 
        key="editable_grafico_nec_enunciado"
    )
    st.text_area(
        "Datos del Gráfico (Enunciado)", 
        key="editable_grafico_json_enunciado", 
        height=150
    )
    
    # --- PREVISUALIZACIÓN (Enunciado) ---
    if st.session_state.editable_grafico_nec_enunciado == "SÍ" and GRAFICOS_DISPONIBLES:
        with st.expander("Previsualizar Gráfico del Enunciado"):
            try:
                json_data = json.loads(st.session_state.editable_grafico_json_enunciado)
                if json_data and isinstance(json_data, list):
                    spec = json_data[0] # Tomar el primer gráfico de la lista
                    
                    # --- LLAMADA DIRECTA AL RENDERIZADOR ---
                    buffer_imagen = crear_grafico(
                        tipo_grafico=spec.get("tipo_elemento"),
                        datos=spec.get("datos", {}),
                        configuracion=spec.get("configuracion", {})
                    )
                    if buffer_imagen:
                        st.image(buffer_imagen, caption="Previsualización")
                    else:
                        st.error("No se pudo renderizar el gráfico. Revisa el JSON.")

            except json.JSONDecodeError:
                st.error("Error en el formato JSON del gráfico del enunciado.")
            except Exception as e:
                st.error(f"Error al intentar renderizar el gráfico: {e}")

    
    # --- OPCIONES Y SUS GRÁFICOS ---
    st.subheader("Opciones")
    
    for letra in ["A", "B", "C", "D"]:
        st.markdown(f"--- \n**Opción {letra}**")
        st.text_input(f"Texto Opción {letra}", key=f"editable_opcion_{letra.lower()}_texto")
        st.selectbox(
            f"¿Gráfico en Opción {letra}?", 
            options=["NO", "SÍ"], 
            key=f"editable_opcion_{letra.lower()}_grafico_nec"
        )
        st.text_area(
            f"Datos Gráfico Opción {letra} (JSON)", 
            key=f"editable_opcion_{letra.lower()}_grafico_json", 
            height=100
        )
        
        # --- PREVISUALIZACIÓN (Opciones) ---
        if st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"] == "SÍ" and GRAFICOS_DISPONIBLES:
            with st.expander(f"Previsualizar Gráfico de Opción {letra}"):
                try:
                    json_data = json.loads(st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"])
                    if json_data and isinstance(json_data, list):
                        spec = json_data[0]
                        
                        # --- LLAMADA DIRECTA AL RENDERIZADOR ---
                        buffer_imagen = crear_grafico(
                            tipo_grafico=spec.get("tipo_elemento"),
                            datos=spec.get("datos", {}),
                            configuracion=spec.get("configuracion", {})
                        )
                        if buffer_imagen:
                            st.image(buffer_imagen, caption="Previsualización")
                        else:
                            st.error("No se pudo renderizar el gráfico. Revisa el JSON.")
                except json.JSONDecodeError:
                    st.error(f"Error en el formato JSON del gráfico de la Opción {letra}.")
                except Exception as e:
                    st.error(f"Error al intentar renderizar el gráfico: {e}")
        
    st.subheader("Clave")
    st.text_input("Clave (Respuesta Correcta)", key="editable_clave")

    st.subheader("Justificaciones")
    st.text_area("Justificación Clave", key="editable_just_clave", height=100)
    st.text_area("Justificación A", key="editable_just_a", height=100)
    st.text_area("Justificación B", key="editable_just_b", height=100)
    st.text_area("Justificación C", key="editable_just_c", height=100)
    st.text_area("Justificación D", key="editable_just_d", height=100)

    # --- SECCIÓN DE DESCARGA ---
    st.divider()
    st.header("4. Descargar Resultados")
    
    # --- LÓGICA DE RE-ENSAMBLE (ACTUALIZADA) ---
    datos_editados = {
        "pregunta_espejo": st.session_state.editable_pregunta,
        "clave": st.session_state.editable_clave,
        "justificacion_clave": st.session_state.editable_just_clave,
        "grafico_necesario_enunciado": st.session_state.editable_grafico_nec_enunciado,
        "opciones": {},
        "justificaciones_distractores": [
            {"opcion": "A", "justificacion": st.session_state.editable_just_a},
            {"opcion": "B", "justificacion": st.session_state.editable_just_b},
            {"opcion": "C", "justificacion": st.session_state.editable_just_c},
            {"opcion": "D", "justificacion": st.session_state.editable_just_d},
        ]
    }
    
    # Re-ensamble del gráfico del enunciado
    try:
        datos_editados["descripcion_grafico_enunciado"] = json.loads(st.session_state.editable_grafico_json_enunciado)
    except json.JSONDecodeError:
        st.error("El JSON del gráfico del enunciado tiene un error de formato, se guardará como texto.")
        datos_editados["descripcion_grafico_enunciado"] = st.session_state.editable_grafico_json_enunciado
    
    # Re-ensamble de las opciones (A, B, C, D)
    for letra in ["A", "B", "C", "D"]:
        opcion_data = {
            "texto": st.session_state[f"editable_opcion_{letra.lower()}_texto"],
            "grafico_necesario": st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"]
        }
        try:
            opcion_data["descripcion_grafico"] = json.loads(st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"])
        except json.JSONDecodeError:
            opcion_data["descripcion_grafico"] = st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"]
            st.error(f"El JSON del gráfico de la Opción {letra} tiene un error, se guardará como texto.")
        
        datos_editados["opciones"][letra] = opcion_data

    
    col_word, col_excel = st.columns(2)
    
    with col_word:
        archivo_word = crear_word(datos_editados)
        st.download_button(
            label="Descargar en Word (.docx)",
            data=archivo_word,
            file_name="item_espejo_auditado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col_excel:
        archivo_excel = crear_excel(datos_editados)
        st.download_button(
            label="Descargar en Excel (.xlsx)",
            data=archivo_excel,
            file_name="item_espejo_auditado.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
